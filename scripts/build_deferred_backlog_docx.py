from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUTPUT = Path("docs/delivery/backlog-sau-module-1.docx")

NAVY = "0B2545"
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
MUTED = "5D6673"
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
WHITE = "FFFFFF"
GREEN = "1F6E43"
AMBER = "8A6500"
RED = "9B1C1C"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths):
    total = sum(widths)
    table.autofit = False
    tbl_pr = table._tbl.tblPr

    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(total))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")

    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for index, cell in enumerate(row.cells):
            width = widths[index]
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_run(run, size=11, color=None, bold=None, italic=None, font="Calibri"):
    run.font.name = font
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), font)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), font)
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def style_paragraph(paragraph, before=0, after=6, line=1.25):
    paragraph.paragraph_format.space_before = Pt(before)
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = line


def add_numbering_definition(document, fmt, text, left=540, hanging=270):
    numbering = document.part.numbering_part.element
    abstract_ids = [
        int(node.get(qn("w:abstractNumId")))
        for node in numbering.findall(qn("w:abstractNum"))
    ]
    num_ids = [
        int(node.get(qn("w:numId")))
        for node in numbering.findall(qn("w:num"))
    ]
    abstract_id = max(abstract_ids, default=0) + 1
    num_id = max(num_ids, default=0) + 1

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)

    level = OxmlElement("w:lvl")
    level.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    level.append(start)
    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), fmt)
    level.append(num_fmt)
    lvl_text = OxmlElement("w:lvlText")
    lvl_text.set(qn("w:val"), text)
    level.append(lvl_text)
    lvl_jc = OxmlElement("w:lvlJc")
    lvl_jc.set(qn("w:val"), "left")
    level.append(lvl_jc)
    p_pr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), str(left))
    tabs.append(tab)
    p_pr.append(tabs)
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), str(left))
    ind.set(qn("w:hanging"), str(hanging))
    p_pr.append(ind)
    level.append(p_pr)
    abstract.append(level)
    numbering.append(abstract)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    numbering.append(num)
    return num_id


def apply_numbering(paragraph, num_id):
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = p_pr.find(qn("w:numPr"))
    if num_pr is None:
        num_pr = OxmlElement("w:numPr")
        p_pr.append(num_pr)
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num_id_node = OxmlElement("w:numId")
    num_id_node.set(qn("w:val"), str(num_id))
    num_pr.append(ilvl)
    num_pr.append(num_id_node)


def add_bullet(document, text, bullet_id, bold_prefix=None):
    paragraph = document.add_paragraph()
    apply_numbering(paragraph, bullet_id)
    style_paragraph(paragraph, after=4)
    if bold_prefix and text.startswith(bold_prefix):
        first = paragraph.add_run(bold_prefix)
        set_run(first, bold=True)
        rest = paragraph.add_run(text[len(bold_prefix):])
        set_run(rest)
    else:
        set_run(paragraph.add_run(text))
    return paragraph


def add_numbered(document, text, number_id):
    paragraph = document.add_paragraph()
    apply_numbering(paragraph, number_id)
    style_paragraph(paragraph, after=4)
    set_run(paragraph.add_run(text))
    return paragraph


def add_heading(document, text, level=1):
    paragraph = document.add_paragraph(style=f"Heading {level}")
    paragraph.add_run(text)
    return paragraph


def add_body(document, text, bold_lead=None):
    paragraph = document.add_paragraph()
    style_paragraph(paragraph)
    if bold_lead and text.startswith(bold_lead):
        lead = paragraph.add_run(bold_lead)
        set_run(lead, bold=True)
        set_run(paragraph.add_run(text[len(bold_lead):]))
    else:
        set_run(paragraph.add_run(text))
    return paragraph


def add_status_table(document, headers, rows, widths):
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_geometry(table, widths)
    repeat_table_header(table.rows[0])
    for index, header in enumerate(headers):
        cell = table.rows[0].cells[index]
        set_cell_shading(cell, LIGHT_BLUE)
        paragraph = cell.paragraphs[0]
        style_paragraph(paragraph, after=0, line=1.1)
        set_run(paragraph.add_run(header), size=10, color=NAVY, bold=True)

    for row_data in rows:
        row = table.add_row()
        for index, value in enumerate(row_data):
            cell = row.cells[index]
            if index == len(row_data) - 1:
                status = value.lower()
                if "hoàn thành" in status or "đạt" in status:
                    set_cell_shading(cell, "E8F3EC")
                elif "hoãn" in status or "chưa" in status:
                    set_cell_shading(cell, "FFF4D6")
            paragraph = cell.paragraphs[0]
            style_paragraph(paragraph, after=0, line=1.1)
            color = NAVY if index == 0 else None
            set_run(paragraph.add_run(value), size=9.5, color=color, bold=index == 0)
        set_table_geometry(table, widths)
    document.add_paragraph().paragraph_format.space_after = Pt(2)
    return table


def add_callout(document, label, text, fill=LIGHT_GRAY, color=NAVY):
    paragraph = document.add_paragraph()
    style_paragraph(paragraph, before=4, after=8, line=1.15)
    paragraph.paragraph_format.left_indent = Inches(0.12)
    paragraph.paragraph_format.right_indent = Inches(0.08)
    p_pr = paragraph._p.get_or_add_pPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    p_pr.append(shading)
    borders = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "18")
    left.set(qn("w:space"), "8")
    left.set(qn("w:color"), color)
    borders.append(left)
    p_pr.append(borders)
    set_run(paragraph.add_run(f"{label}: "), bold=True, color=color)
    set_run(paragraph.add_run(text), color=color)


def configure_styles(document):
    normal = document.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    tokens = {
        "Title": (25, NAVY, 0, 5),
        "Subtitle": (12.5, MUTED, 0, 16),
        "Heading 1": (16, BLUE, 18, 10),
        "Heading 2": (13, BLUE, 14, 7),
        "Heading 3": (12, DARK_BLUE, 10, 5),
    }
    for style_name, (size, color, before, after) in tokens.items():
        style = document.styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = style_name != "Subtitle"
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.1
        style.paragraph_format.keep_with_next = True


def add_footer(section):
    footer = section.footer
    paragraph = footer.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    style_paragraph(paragraph, after=0, line=1)
    set_run(paragraph.add_run("Backlog sau Module 1  |  Trang "), size=9, color=MUTED)
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_begin, instr, fld_end])
    set_run(run, size=9, color=MUTED)


def add_header(section):
    header = section.header
    paragraph = header.paragraphs[0]
    style_paragraph(paragraph, after=0, line=1)
    set_run(paragraph.add_run("SEEDING STRATEGY TOOL"), size=8.5, color=MUTED, bold=True)


def new_page(document):
    document.add_page_break()


def build_document():
    document = Document()
    section = document.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    add_header(section)
    add_footer(section)
    configure_styles(document)

    bullet_id = add_numbering_definition(document, "bullet", "•")
    number_id = add_numbering_definition(document, "decimal", "%1.")

    # First-page memo masthead.
    kicker = document.add_paragraph()
    style_paragraph(kicker, before=10, after=3, line=1)
    set_run(kicker.add_run("TÀI LIỆU KIỂM SOÁT PHẠM VI"), size=9, color=BLUE, bold=True)

    title = document.add_paragraph(style="Title")
    title.add_run("Backlog sau khi hoàn thành Module 1")

    subtitle = document.add_paragraph(style="Subtitle")
    subtitle.add_run(
        "Business & AnalysisSession - danh sách công việc tạm hoãn và kế hoạch tiếp tục"
    )

    metadata = [
        ("Dự án", "Seeding Strategy Tool"),
        ("Phiên bản", "1.0"),
        ("Ngày chốt", "28/07/2026"),
        ("Trạng thái", "Module 1 sẵn sàng thử nghiệm API"),
        ("Phạm vi tạm hoãn", "Authentication, frontend và Giai đoạn 2-5"),
    ]
    for label, value in metadata:
        paragraph = document.add_paragraph()
        style_paragraph(paragraph, after=2, line=1.1)
        set_run(paragraph.add_run(f"{label}: "), bold=True, color=NAVY)
        set_run(paragraph.add_run(value))

    add_callout(
        document,
        "Quyết định phạm vi",
        "Trong đợt thử Module 1, hệ thống sử dụng ba header tạm thời để mô phỏng người dùng và organization. Đăng nhập, JWT và kiểm tra membership được chủ động hoãn; không được xem là đã hoàn thành.",
        fill="EAF2F8",
    )

    add_heading(document, "1. Trạng thái Module 1", 1)
    add_body(
        document,
        "Module 1 trong tài liệu này gồm hai bounded context: Business và AnalysisSession. Phần backend đã có migration, API contract, Swagger, validation, transaction, unit test và E2E test với PostgreSQL thật."
    )
    add_status_table(
        document,
        ["Hạng mục", "Kết quả hiện tại", "Trạng thái"],
        [
            ("Business", "Tạo, xem, lọc, cập nhật, deactivate và restore.", "Hoàn thành"),
            (
                "AnalysisSession",
                "Tạo DRAFT, cập nhật phạm vi, start, complete, archive và theo dõi trạng thái.",
                "Hoàn thành",
            ),
            (
                "An toàn dữ liệu",
                "Tenant scope, snapshot Business, optimistic update và transaction rollback.",
                "Đạt",
            ),
            (
                "API & kiểm thử",
                "Swagger hoạt động; 46 unit/contract test và 9 E2E test đạt.",
                "Đạt",
            ),
            (
                "Authentication",
                "Đang dùng x-organization-id, x-user-id và x-user-role.",
                "Hoãn có chủ đích",
            ),
            ("Frontend", "Chưa triển khai trong phạm vi hiện tại.", "Hoãn có chủ đích"),
        ],
        [2100, 5260, 2000],
    )

    add_heading(document, "2. Điều kiện chấp nhận Module 1", 1)
    for item in [
        "Swagger UI mở được tại /api/docs và OpenAPI JSON tại /api/docs-json.",
        "Có thể tạo và cập nhật Business đang hoạt động.",
        "Có thể tạo AnalysisSession DRAFT cho đúng Business và organization.",
        "Khi start session, hồ sơ Business được đóng băng trong businessSnapshot.",
        "Không thể sửa phạm vi session sau khi rời DRAFT.",
        "Deactivate Business tự archive DRAFT nhưng rollback nếu còn session đang chạy.",
        "Không truy cập chéo dữ liệu giữa hai organization.",
        "Prisma validate, lint, typecheck, unit test, E2E test và build đều đạt.",
    ]:
        add_bullet(document, item, bullet_id)

    add_callout(
        document,
        "Lưu ý thử nghiệm",
        "Ba header tạm chỉ phù hợp cho local/demo nội bộ. Không dùng cơ chế này để triển khai production hoặc cấp quyền cho người dùng thật.",
        fill="FFF4D6",
        color=AMBER,
    )

    add_heading(document, "2.1 Cấu hình thử nhanh bằng Swagger", 2)
    for item in [
        "Chạy migration: pnpm --filter @seeding/api db:migrate:deploy.",
        "Tạo tenant demo: pnpm --filter @seeding/api db:seed.",
        "Khởi động API: pnpm --filter @seeding/api dev.",
        "Mở http://localhost:3001/api/docs.",
        "Dùng x-organization-id=org_demo, x-user-id=user_demo_admin, x-user-role=ORG_ADMIN.",
        "Khi thử quyền Analyst, đổi x-user-id=user_demo_analyst và x-user-role=ANALYST.",
    ]:
        add_numbered(document, item, number_id)

    new_page(document)
    add_heading(document, "3. Công việc nền tảng tạm hoãn", 1)
    add_heading(document, "3.1 Authentication và phiên đăng nhập", 2)
    for item in [
        "Xây dựng API đăng nhập, access token, refresh token, logout và thu hồi token.",
        "Hash mật khẩu an toàn; chính sách mật khẩu, khóa tài khoản và xử lý token hết hạn.",
        "Thay TemporaryRequestContextMiddleware bằng JWT Guard.",
        "Không nhận userId hoặc role từ client; lấy danh tính từ token đã xác minh.",
        "Bổ sung E2E cho đăng nhập đúng/sai, refresh, logout và token bị thu hồi.",
    ]:
        add_bullet(document, item, bullet_id)

    add_heading(document, "3.2 Organization và Membership", 2)
    for item in [
        "API tạo, xem, cập nhật và chọn Organization hiện tại.",
        "API quản lý OrganizationMember và trạng thái thành viên.",
        "Đọc role từ OrganizationMember trong database.",
        "Chặn người dùng không thuộc organization hoặc membership đã inactive.",
        "Bổ sung phân quyền chi tiết theo Business cho trường hợp Analyst được cấp quyền.",
    ]:
        add_bullet(document, item, bullet_id)

    add_heading(document, "3.3 Hardening trước production", 2)
    for item in [
        "Cấu hình CORS theo domain; bổ sung Helmet, rate limiting và giới hạn request body.",
        "Structured logging cho request thành công/thất bại và truyền requestId xuyên suốt.",
        "Quản lý secret theo môi trường; không dùng secret mẫu.",
        "Đưa migration, unit test, contract test và E2E vào CI.",
        "Bổ sung test đồng thời thực sự cho start session và deactivate Business.",
        "Cân nhắc optimistic version/ETag để tránh hai người ghi đè Business.",
    ]:
        add_bullet(document, item, bullet_id)

    add_heading(document, "3.4 Quyết định nghiệp vụ cần xác nhận", 2)
    add_status_table(
        document,
        ["Quyết định", "Nội dung cần chốt", "Thời điểm"],
        [
            (
                "Complete sớm",
                "Có tiếp tục cho DATA_COLLECTION chuyển thẳng COMPLETED khi Giai đoạn 2-5 hoạt động hay không.",
                "Trước GĐ2",
            ),
            (
                "Auto-archive DRAFT",
                "Xác nhận deactivate Business được phép tự archive toàn bộ session DRAFT.",
                "Sau thử M1",
            ),
            (
                "Audit trạng thái",
                "Có cần bảng lịch sử người chuyển trạng thái session và lý do hay chỉ lưu timestamp.",
                "Trước production",
            ),
            (
                "Lưu trữ dữ liệu",
                "Chính sách retention cho file import, feedback, raw AI response và bản lỗi.",
                "Trước GĐ2",
            ),
        ],
        [1900, 5700, 1760],
    )

    new_page(document)
    add_heading(document, "4. Giai đoạn 2 - Nhập dữ liệu", 1)
    add_body(
        document,
        "Schema nền đã có DataSource, ImportBatch và CustomerFeedback, nhưng chưa có module NestJS, repository, service, controller hoặc API contract cho luồng nhập dữ liệu."
    )

    add_heading(document, "4.1 DataSource và feedback thủ công", 2)
    for item in [
        "Tạo module DataSource và CustomerFeedback theo cùng ranh giới tenant/business/session.",
        "API tạo, xem, sửa và archive feedback nhập thủ công.",
        "Validate nội dung, rating, thời gian, URL nguồn và metadata.",
        "Lưu nguyên bản rawContent và dữ liệu đã chuẩn hóa ở trường riêng.",
        "Chỉ cho phép nhập dữ liệu khi session ở trạng thái DATA_COLLECTION.",
    ]:
        add_bullet(document, item, bullet_id)

    add_heading(document, "4.2 Upload Excel/CSV và ImportBatch", 2)
    for item in [
        "Upload file; kiểm tra extension, MIME, dung lượng, file rỗng và encoding.",
        "Lưu file vào local storage hoặc object storage theo cấu hình.",
        "Đọc sheet/header, cho phép chọn sheet và xem trước dữ liệu.",
        "Cung cấp API mapping cột chuẩn với cột trong file.",
        "Lưu mapping, trạng thái và thống kê từng ImportBatch.",
        "Hỗ trợ idempotency để không import trùng khi client retry.",
    ]:
        add_bullet(document, item, bullet_id)

    add_heading(document, "4.3 Validation và file lỗi", 2)
    for item in [
        "Validate từng dòng và ghi mã lỗi, tên cột, giá trị gốc, thông báo.",
        "Chỉ lưu feedback hợp lệ; không làm mất dữ liệu gốc.",
        "Xuất CSV/XLSX chứa dòng lỗi và cột giải thích.",
        "Theo dõi tổng dòng, hợp lệ, lỗi, đã import và bỏ qua.",
        "E2E cho file hợp lệ, sai định dạng, thiếu header, import lại và file lớn.",
    ]:
        add_bullet(document, item, bullet_id)

    add_heading(document, "4.4 Điều kiện hoàn thành Giai đoạn 2", 2)
    for item in [
        "Nhập feedback thủ công và bằng file đều hoạt động.",
        "Mapping cột có thể lưu và tái sử dụng khi phù hợp.",
        "Dữ liệu hợp lệ được gắn đúng DataSource, Business và AnalysisSession.",
        "Người dùng tải được file lỗi; số liệu ImportBatch khớp dữ liệu thực.",
    ]:
        add_bullet(document, item, bullet_id)

    new_page(document)
    add_heading(document, "5. Giai đoạn 3 - Xử lý dữ liệu và AI", 1)
    add_body(
        document,
        "Schema đã có ProcessingJob và FeedbackAnalysis dạng versioned run, nhưng queue, worker, AI adapter và API theo dõi tiến độ chưa được triển khai."
    )

    add_heading(document, "5.1 Hạ tầng xử lý nền", 2)
    for item in [
        "Cấu hình Redis, BullMQ queue, worker, concurrency, timeout và retry policy.",
        "Thiết kế idempotency key và cơ chế chống xử lý cùng một batch nhiều lần.",
        "Ghi ProcessingJob theo session, DataSource hoặc ImportBatch bằng cột riêng.",
        "API xem tiến độ, tổng item, item thành công/thất bại và lỗi gần nhất.",
        "Cơ chế retry thủ công, cancel và phục hồi job bị treo.",
    ]:
        add_bullet(document, item, bullet_id)

    add_heading(document, "5.2 Chuẩn hóa và phát hiện trùng", 2)
    for item in [
        "Chuẩn hóa Unicode, khoảng trắng, ngôn ngữ, ngày, rating và URL.",
        "Giữ raw content bất biến; tạo normalizedContent và contentHash.",
        "Phát hiện trùng theo externalId, contentHash và rule bổ sung.",
        "Đánh dấu duplicateOfId; không xóa vật lý feedback trùng.",
        "Test dữ liệu không dấu, sai chính tả, emoji, text rất ngắn và nội dung rỗng.",
    ]:
        add_bullet(document, item, bullet_id)

    add_heading(document, "5.3 AI FeedbackAnalysis", 2)
    for item in [
        "Xây AI Provider Adapter để thay đổi OpenAI/Gemini mà không đổi use case.",
        "Định nghĩa và validate JSON schema cho sentiment, topic, pain point, question, priority, confidence và evidence.",
        "Lưu aiModel, promptVersion, rawResponse, runNo và errorMessage.",
        "Không ghi đè lần chạy cũ; cho phép so sánh và chọn kết quả hiện hành.",
        "Theo dõi token/cost, rate limit, timeout, retry và lỗi output không đúng schema.",
        "Chuẩn bị bộ dữ liệu đánh giá AI có expected result.",
    ]:
        add_bullet(document, item, bullet_id)

    add_heading(document, "5.4 Điều kiện hoàn thành Giai đoạn 3", 2)
    for item in [
        "Toàn bộ feedback hợp lệ được chuẩn hóa và phân loại duplicate.",
        "AI analysis có schema hợp lệ, version và khả năng retry.",
        "Tiến độ job phản ánh số liệu thật; lỗi không làm mất kết quả đã hoàn thành.",
        "Session chỉ chuyển bước khi các job bắt buộc hoàn tất.",
    ]:
        add_bullet(document, item, bullet_id)

    new_page(document)
    add_heading(document, "6. Giai đoạn 4 - Insight", 1)
    add_body(
        document,
        "Schema đã có Insight, InsightEvidence và InsightReviewLog. Chưa có module tổng hợp, API review, policy reviewer hoặc luồng trạng thái hoàn chỉnh."
    )

    add_heading(document, "6.1 Sinh và tổng hợp Insight", 2)
    for item in [
        "Nhóm feedback theo chủ đề/pain point và tạo insight có priority, confidence, origin.",
        "Mỗi insight phải liên kết evidence cụ thể bằng InsightEvidence.",
        "Phân biệt insight observed, inferred và assumed.",
        "Không tạo insight nếu không đủ evidence theo ngưỡng đã chốt.",
        "Lưu model, prompt version và nguồn dữ liệu dùng để tổng hợp.",
    ]:
        add_bullet(document, item, bullet_id)

    add_heading(document, "6.2 Review và audit", 2)
    for item in [
        "API sửa title/description/priority khi trạng thái cho phép.",
        "Submit review, approve, reject, yêu cầu re-analysis, archive.",
        "Ghi toàn bộ hành động vào InsightReviewLog; không chỉ ghi đè reviewedBy.",
        "Hiển thị feedback bằng chứng và đường dẫn về DataSource.",
        "Chính sách merge/split tạo insight mới qua parentInsightId thay vì sửa lịch sử.",
        "Chặn chỉnh sửa insight đang được strategy locked tham chiếu.",
    ]:
        add_bullet(document, item, bullet_id)

    add_heading(document, "6.3 Gate sang Strategy", 2)
    for item in [
        "Chỉ Insight APPROVED được dùng tạo Strategy.",
        "Rejected/archived insight không xuất hiện trong input strategy.",
        "Khi insight thay đổi sau duyệt, phải quay lại review hoặc tạo phiên bản logic mới.",
        "E2E bao phủ approve/reject/re-submit, evidence, merge/split và tenant isolation.",
    ]:
        add_bullet(document, item, bullet_id)

    new_page(document)
    add_heading(document, "7. Giai đoạn 5 - Strategy", 1)
    add_body(
        document,
        "Schema đã có Strategy, StrategyVersion và StrategyInsight cùng insightSnapshot. Chưa có service tạo chiến lược, version workflow, approval/lock và output contract cho module nội dung."
    )

    add_heading(document, "7.1 Tạo chiến lược bằng AI", 2)
    for item in [
        "Chỉ lấy Business snapshot và danh sách Insight APPROVED.",
        "Định nghĩa JSON schema đầu ra: mục tiêu, đối tượng, thông điệp, kênh, trụ cột nội dung và guardrail.",
        "Lưu model, promptVersion và lỗi sinh chiến lược.",
        "Lưu StrategyInsight kèm insightSnapshot để đóng băng evidence đầu vào.",
        "Không sinh hai bản trùng khi client retry.",
    ]:
        add_bullet(document, item, bullet_id)

    add_heading(document, "7.2 Versioning, chỉnh sửa và duyệt", 2)
    for item in [
        "Tạo versionNo tuần tự trong transaction.",
        "Manager chỉnh sửa bản DRAFT; không sửa trực tiếp bản LOCKED.",
        "Submit, request revision, approve, lock, supersede và archive.",
        "Ghi createdBy, approvedBy, approvedAt, lockedAt và review comment.",
        "Tạo version mới từ bản locked thay vì mở khóa và ghi đè.",
    ]:
        add_bullet(document, item, bullet_id)

    add_heading(document, "7.3 Truy vết và dữ liệu đầu ra", 2)
    for item in [
        "Từ mỗi phần chiến lược truy ngược được Insight và feedback bằng chứng.",
        "Output contract ổn định cho module tạo nội dung.",
        "Chỉ phát hành version APPROVED/LOCKED sang module sau.",
        "Export JSON có version, timestamp, Business snapshot và insight references.",
        "E2E kiểm tra versioning, approval, lock, immutable snapshot và tenant isolation.",
    ]:
        add_bullet(document, item, bullet_id)

    new_page(document)
    add_heading(document, "8. Thứ tự thực hiện sau Module 1", 1)
    sequence = [
        "Cho người dùng thử Module 1 bằng Swagger; ghi nhận lỗi nghiệp vụ và chốt ba quyết định ở mục 3.4.",
        "Hoàn thiện Authentication, Organization Membership và RBAC thật.",
        "Triển khai Giai đoạn 2: nhập thủ công trước, sau đó Excel/CSV.",
        "Triển khai Giai đoạn 3: normalization/dedup trước, AI analysis sau.",
        "Triển khai Giai đoạn 4: Insight generation, evidence và review workflow.",
        "Triển khai Giai đoạn 5: Strategy generation, versioning, approval và output.",
        "Cuối mỗi giai đoạn: cập nhật contract, Swagger, migration, unit test, E2E và tài liệu này.",
    ]
    for item in sequence:
        add_numbered(document, item, number_id)

    add_heading(document, "9. Ma trận ưu tiên backlog", 1)
    add_status_table(
        document,
        ["Nhóm", "Mục tiêu", "Ưu tiên", "Phụ thuộc"],
        [
            ("Auth/RBAC", "Danh tính và quyền thật", "P1", "Module 1 ổn định"),
            ("Giai đoạn 2", "Có dữ liệu feedback sạch đầu vào", "P1", "Session DATA_COLLECTION"),
            ("Giai đoạn 3", "Xử lý và phân tích AI có kiểm soát", "P1", "Giai đoạn 2"),
            ("Giai đoạn 4", "Insight có evidence và review", "P1", "Giai đoạn 3"),
            ("Giai đoạn 5", "Strategy có version và approval", "P1", "Giai đoạn 4"),
            ("Production hardening", "Bảo mật, logging và CI", "P1 trước production", "Auth + tất cả module"),
            ("Frontend", "UI theo contract đã ổn định", "Sau API tương ứng", "Theo từng giai đoạn"),
        ],
        [1700, 3500, 1800, 2360],
    )

    add_heading(document, "10. Quy tắc cập nhật tài liệu", 1)
    for item in [
        "Mỗi công việc chỉ chuyển sang Hoàn thành khi code, migration, contract và test tương ứng đều đạt.",
        "Nếu thay đổi quyết định nghiệp vụ, ghi ngày chốt và ảnh hưởng tới migration/API.",
        "Không xóa công việc hoãn; chuyển sang Hoàn thành và ghi phiên bản bàn giao.",
        "Lỗi P1 phải sửa trước khi đóng giai đoạn; P2/P3 được ghi rõ vào backlog.",
        "Tài liệu này là backlog phạm vi; source code và test vẫn là bằng chứng thực thi.",
    ]:
        add_bullet(document, item, bullet_id)

    add_callout(
        document,
        "Điểm tiếp tục",
        "Sau khi Module 1 được thử nghiệm và các quyết định nghiệp vụ được xác nhận, bắt đầu từ Authentication/RBAC hoặc Giai đoạn 2 theo quyết định sản phẩm; không cần thiết kế lại Business và AnalysisSession.",
        fill="E8F3EC",
        color=GREEN,
    )

    properties = document.core_properties
    properties.title = "Backlog sau khi hoàn thành Module 1"
    properties.subject = "Công việc tạm hoãn và kế hoạch Giai đoạn 2-5"
    properties.author = "Seeding Strategy Tool Team"
    properties.keywords = "Module 1, Business, AnalysisSession, backlog, roadmap"

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    document.save(OUTPUT)
    print(OUTPUT.resolve())


if __name__ == "__main__":
    build_document()
